#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Phase 1 Analysis: inspect the structure of Part III_word.docx.

Prints:
  - All unique paragraph styles found
  - First 200 chars of each paragraph with its style name, bold flag, font size
  - Footnote/endnote references if any
  - A summary of structural patterns

Run from repo root:
    python scripts/analyze_part3_docx.py
"""

import io
import os
import sys
from pathlib import Path

# Force UTF-8 output on Windows
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

DOCX_PATH = Path("assets/db/Part III_word.docx")


def ensure_docx():
    try:
        import docx  # noqa: F401
        return True
    except ImportError:
        return False


def analyze(docx_path: Path):
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(str(docx_path))

    styles_seen = {}
    print("=" * 80)
    print(f"Document: {docx_path}")
    print(f"Total paragraphs: {len(doc.paragraphs)}")

    # --- Count unique styles ---
    for p in doc.paragraphs:
        s = p.style.name if p.style else "None"
        styles_seen[s] = styles_seen.get(s, 0) + 1

    print("\nUnique paragraph styles:")
    for style, count in sorted(styles_seen.items(), key=lambda x: -x[1]):
        print(f"  {count:>5}x  {style!r}")

    # --- Print every paragraph with metadata ---
    print("\n" + "=" * 80)
    print("PARAGRAPH DUMP (index | style | bold | font_size | first 200 chars)")
    print("=" * 80)
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        style_name = p.style.name if p.style else "None"

        # Check if any run is bold
        is_bold = any(run.bold for run in p.runs if run.text.strip())
        # First non-None font size
        font_size = None
        for run in p.runs:
            if run.font.size:
                font_size = run.font.size.pt
                break

        preview = text[:200].replace("\n", "\\n")
        print(f"[{i:04d}] style={style_name!r:30s} bold={str(is_bold):<5} "
              f"sz={str(font_size) if font_size else '?':>5}pt  |  {preview}")

    # --- Check for Word footnotes/endnotes ---
    print("\n" + "=" * 80)
    print("WORD FOOTNOTES (docx built-in):")
    fn_part = None
    try:
        fn_part = doc.part.footnotes_part
    except Exception:
        pass
    if fn_part is not None:
        fn_doc = fn_part._element
        fns = fn_doc.findall(f".//{qn('w:footnote')}")
        print(f"  Found {len(fns)} Word footnote(s)")
        for fn in fns[:5]:
            fn_id = fn.get(qn('w:id'))
            texts = [t.text or '' for t in fn.iter(qn('w:t'))]
            print(f"  id={fn_id}: {''.join(texts)[:120]!r}")
        if len(fns) > 5:
            print(f"  ... and {len(fns) - 5} more")
    else:
        print("  No built-in Word footnotes found")

    # --- Check for Word endnotes ---
    print("\nWORD ENDNOTES (docx built-in):")
    en_part = None
    try:
        en_part = doc.part.endnotes_part
    except Exception:
        pass
    if en_part is not None:
        en_doc = en_part._element
        ens = en_doc.findall(f".//{qn('w:endnote')}")
        print(f"  Found {len(ens)} Word endnote(s)")
        for en in ens[:5]:
            en_id = en.get(qn('w:id'))
            texts = [t.text or '' for t in en.iter(qn('w:t'))]
            print(f"  id={en_id}: {''.join(texts)[:120]!r}")
    else:
        print("  No built-in Word endnotes found")

    # --- Look for inline footnote ref patterns in text ---
    import re
    print("\n" + "=" * 80)
    print("INLINE REF PATTERNS in paragraph text:")
    patterns = {
        r'\[\^?\d+\]': '[^N] or [N]',
        r'\(\d+\)': '(N)',
        r'^\d+\s': 'Starts with digit+space (footnote text?)',
        r'\d+$': 'Ends with digit',
    }
    pattern_hits = {k: [] for k in patterns}

    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        for pat, desc in patterns.items():
            if re.search(pat, text):
                pattern_hits[pat].append((i, text[:100]))

    for pat, desc in patterns.items():
        hits = pattern_hits[pat]
        print(f"\n  Pattern {desc!r} — {len(hits)} hits (first 5):")
        for idx, txt in hits[:5]:
            print(f"    [{idx:04d}] {txt!r}")

    # --- Try to detect surah boundaries ---
    print("\n" + "=" * 80)
    print("POTENTIAL SURAH HEADINGS (bold or heading style, short lines ≤50 chars):")
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text or len(text) > 80:
            continue
        style_name = p.style.name if p.style else "None"
        is_bold = any(run.bold for run in p.runs if run.text.strip())
        is_heading = 'heading' in style_name.lower() or 'title' in style_name.lower()
        if is_bold or is_heading:
            print(f"  [{i:04d}] style={style_name!r:30s} bold={is_bold}  |  {text!r}")

    # --- Print first 50 paragraphs for context ---
    print("\n" + "=" * 80)
    print("FIRST 60 NON-EMPTY PARAGRAPHS (with full style info):")
    count = 0
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        style_name = p.style.name if p.style else "None"
        is_bold = any(run.bold for run in p.runs if run.text.strip())
        print(f"  [{i:04d}] {style_name!r:30s} bold={is_bold}  |  {text[:150]!r}")
        count += 1
        if count >= 60:
            break

    # --- Print last 30 paragraphs ---
    print("\n" + "=" * 80)
    print("LAST 30 NON-EMPTY PARAGRAPHS:")
    all_nonempty = [(i, p) for i, p in enumerate(doc.paragraphs) if p.text.strip()]
    for i, p in all_nonempty[-30:]:
        text = p.text.strip()
        style_name = p.style.name if p.style else "None"
        is_bold = any(run.bold for run in p.runs if run.text.strip())
        print(f"  [{i:04d}] {style_name!r:30s} bold={is_bold}  |  {text[:150]!r}")


def main():
    if not DOCX_PATH.exists():
        print(f"ERROR: {DOCX_PATH} not found. Run from repo root.", file=sys.stderr)
        sys.exit(1)

    if not ensure_docx():
        print("python-docx not installed. Installing...")
        import subprocess
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        print("Installed. Re-importing...")

    analyze(DOCX_PATH)


if __name__ == "__main__":
    main()
